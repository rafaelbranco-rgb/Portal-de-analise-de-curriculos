"""Extração de texto de arquivos de currículo (PDF, DOCX, TXT)."""
from __future__ import annotations

import io
from pathlib import Path

from pypdf import PdfReader
from docx import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class UnsupportedFileError(ValueError):
    pass


def extract_text(filename: str, raw_bytes: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(raw_bytes)
    if ext == ".docx":
        return _extract_docx(raw_bytes)
    if ext == ".txt":
        return raw_bytes.decode("utf-8", errors="ignore")
    raise UnsupportedFileError(
        f"Formato não suportado: {ext}. Use PDF, DOCX ou TXT."
    )


def _extract_pdf(raw_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(raw_bytes))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n".join(parts).strip()


def _extract_docx(raw_bytes: bytes) -> str:
    document = Document(io.BytesIO(raw_bytes))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts).strip()
